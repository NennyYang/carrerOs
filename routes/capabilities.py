from datetime import datetime, timezone
import logging
from io import BytesIO
from pathlib import Path

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile, status
from sqlalchemy import func
from sqlalchemy.exc import IntegrityError
from sqlalchemy.orm import Session

from database import get_db
from llm_service import (
    analyze_project_with_llm,
    analyze_resume_dimensions_with_llm,
    analyze_resume_match_with_llm,
    analyze_resume_with_llm,
    generate_learning_suggestions_with_llm,
)
from models import (
    JobSkill,
    LearningBacklog,
    ProjectCapabilityProfile,
    ProjectDimensionScoreRecord,
    ResumeAnalysisRecord,
    ResumeDimensionScoreRecord,
    ResumeVersion,
    User,
)
from schemas import (
    JobSkillCreate,
    JobSkillResponse,
    LearningBacklogCreate,
    LearningBacklogResponse,
    LearningSuggestionRequest,
    LearningSuggestionResponse,
    ProjectAnalysisRequest,
    ProjectAnalysisResponse,
    ProjectCapabilityProfileResponse,
    ProjectCapabilityProfileUpsert,
    ProjectDimensionScoreRecordResponse,
    ProjectDimensionScoreRequest,
    ResumeAnalysisRequest,
    ResumeAnalysisResponse,
    ResumeDimensionScoreRecordResponse,
    ResumeDimensionsRequest,
    ResumeDimensionsResponse,
    ResumeMatchAnalysisCreate,
    ResumeMatchAnalysisResponse,
    ResumeVersionCreate,
    ResumeVersionResponse,
)

router = APIRouter(prefix="/api/capabilities", tags=["capabilities"])
logger = logging.getLogger("careeros.capabilities")

PDF_SUFFIX = ".pdf"
DOCX_SUFFIX = ".docx"


def normalize_skill_name(skill_name: str) -> str:
    return " ".join(skill_name.strip().lower().split())


def ensure_user_exists(user_id: int, db: Session) -> User:
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="User not found",
        )
    return user


def next_resume_version_no(user_id: int, db: Session) -> int:
    latest_version = (
        db.query(func.max(ResumeVersion.version_no))
        .filter(ResumeVersion.user_id == user_id)
        .scalar()
        or 0
    )
    return latest_version + 1


def save_resume_version(
    user_id: int,
    filename: str,
    content: str,
    db: Session,
) -> ResumeVersion:
    resume = ResumeVersion(
        user_id=user_id,
        version_no=next_resume_version_no(user_id, db),
        filename=filename.strip(),
        content=content.strip(),
    )
    db.add(resume)
    db.commit()
    db.refresh(resume)
    return resume


def extract_pdf_text(file_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as error:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="PDF extraction requires dependency pypdf. Run dependency install first.",
        ) from error

    reader = PdfReader(BytesIO(file_bytes))
    page_texts = [(page.extract_text() or "") for page in reader.pages]
    return "\n\n".join(text.strip() for text in page_texts if text.strip())


def extract_docx_text(file_bytes: bytes) -> str:
    try:
        from docx import Document
    except ImportError as error:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="DOCX extraction requires dependency python-docx. Run dependency install first.",
        ) from error

    document = Document(BytesIO(file_bytes))
    paragraph_texts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_texts = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_texts.append(" | ".join(cells))
    return "\n".join(paragraph_texts + table_texts)


def extract_resume_text(filename: str, file_bytes: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == PDF_SUFFIX:
        text = extract_pdf_text(file_bytes)
    elif suffix == DOCX_SUFFIX:
        text = extract_docx_text(file_bytes)
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Only PDF and DOCX resume files are supported",
        )

    text = text.strip()
    if not text:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Could not extract text from this resume file",
        )
    return text


def score_resume_against_jd(resume_content: str, job_description: str) -> dict:
    resume = resume_content.lower()
    jd = job_description.lower()
    dimensions = [
        ("Technical keyword coverage", ["python", "fastapi", "rag", "agent", "docker", "sql", "api"]),
        ("Project outcome expression", ["evaluation", "result", "improve", "accuracy", "recall", "metric"]),
        ("Engineering delivery", ["deploy", "deployment", "log", "monitor", "test", "postgresql", "interface"]),
        ("Role fit expression", ["business", "requirement", "collaboration", "delivery", "online"]),
    ]
    dimension_scores = []
    suggestions = []
    for name, keywords in dimensions:
        jd_hits = [keyword for keyword in keywords if keyword in jd]
        expected = jd_hits or keywords[:4]
        matched = [keyword for keyword in expected if keyword in resume]
        score = min(96, 52 + int(len(matched) / max(len(expected), 1) * 44))
        missing = [keyword for keyword in expected if keyword not in resume][:3]
        suggestion = (
            "Core signals covered. Add quantified outcomes."
            if not missing
            else "Add evidence for: " + ", ".join(missing)
        )
        dimension_scores.append({"name": name, "score": score, "suggestion": suggestion})
        suggestions.append(suggestion)

    resume_score = round(sum(item["score"] for item in dimension_scores) / len(dimension_scores))
    jd_terms = [
        term
        for term in ["python", "fastapi", "rag", "agent", "docker", "evaluation", "deploy", "api", "data"]
        if term in jd
    ]
    matched_terms = [term for term in jd_terms if term in resume]
    match_score = min(98, 50 + int(len(matched_terms) / max(len(jd_terms), 1) * 48))
    highlights = ["Resume covers " + term for term in matched_terms[:4]] or [
        "Resume version is stored and ready for matching"
    ]
    gaps = [
        "Add a verifiable project bullet for JD term: " + term
        for term in jd_terms
        if term not in resume
    ][:4]
    if not gaps:
        gaps = ["Map matched keywords into project bullets and add measurable outcomes."]
    return {
        "resume_score": resume_score,
        "match_score": match_score,
        "dimension_scores": dimension_scores,
        "resume_suggestions": suggestions,
        "match_highlights": highlights,
        "match_gaps": gaps,
    }


@router.post("/learning-suggestions", response_model=LearningSuggestionResponse)
def generate_learning_suggestions(request: LearningSuggestionRequest):
    logger.info("Learning suggestions API called")
    try:
        result = generate_learning_suggestions_with_llm(request.work_content.strip())
        logger.info(
            "Learning suggestions API completed: model=%s count=%s",
            result.get("model"),
            len(result.get("suggestions", [])),
        )
        return result
    except Exception as error:
        logger.exception("Learning suggestions API failed")
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail=str(error),
        ) from error


@router.post(
    "/job-skills",
    response_model=JobSkillResponse,
    status_code=status.HTTP_201_CREATED,
)
def create_job_skill(skill_data: JobSkillCreate, db: Session = Depends(get_db)):
    position_name = skill_data.position_name.strip()
    skill_name = skill_data.skill_name.strip()
    job_skill = JobSkill(
        position_name=position_name,
        skill_name=skill_name,
        normalized_name=normalize_skill_name(skill_name),
        category=skill_data.category.strip() if skill_data.category else None,
        importance=skill_data.importance,
        description=skill_data.description,
    )
    try:
        db.add(job_skill)
        db.commit()
        db.refresh(job_skill)
    except IntegrityError:
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="This role already has the same skill",
        )
    return job_skill


@router.get("/job-skills", response_model=list[JobSkillResponse])
def list_job_skills(position_name: str | None = None, db: Session = Depends(get_db)):
    query = db.query(JobSkill)
    if position_name:
        query = query.filter(
            func.lower(JobSkill.position_name) == position_name.strip().lower()
        )
    return query.order_by(JobSkill.position_name, JobSkill.skill_name).all()


@router.post("/projects", response_model=ProjectCapabilityProfileResponse)
def save_project_profile(
    project_data: ProjectCapabilityProfileUpsert, db: Session = Depends(get_db)
):
    ensure_user_exists(project_data.user_id, db)

    project_name = project_data.project_name.strip()
    profile = (
        db.query(ProjectCapabilityProfile)
        .filter(
            ProjectCapabilityProfile.user_id == project_data.user_id,
            func.lower(ProjectCapabilityProfile.project_name) == project_name.lower(),
        )
        .first()
    )
    tags = list(
        dict.fromkeys(tag.strip() for tag in project_data.capability_tags if tag.strip())
    )
    if profile is None:
        profile = ProjectCapabilityProfile(
            user_id=project_data.user_id,
            project_name=project_name,
        )
        db.add(profile)

    profile.technology_stack = project_data.technology_stack.strip()
    profile.problem = project_data.problem.strip()
    profile.solution_and_validation = project_data.solution_and_validation.strip()
    profile.capability_tags = tags
    try:
        db.commit()
        db.refresh(profile)
    except IntegrityError:
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="Project profile save conflict",
        )
    return profile


@router.get("/projects", response_model=list[ProjectCapabilityProfileResponse])
def list_project_profiles(user_id: int, db: Session = Depends(get_db)):
    return (
        db.query(ProjectCapabilityProfile)
        .filter(ProjectCapabilityProfile.user_id == user_id)
        .order_by(ProjectCapabilityProfile.updated_at.desc())
        .all()
    )


@router.post(
    "/resumes/upload",
    response_model=ResumeVersionResponse,
    status_code=status.HTTP_201_CREATED,
)
async def upload_resume_file(
    user_id: int = Form(...),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
):
    ensure_user_exists(user_id, db)
    filename = file.filename or ""
    suffix = Path(filename).suffix.lower()
    if suffix not in {PDF_SUFFIX, DOCX_SUFFIX}:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Only PDF and DOCX resume files are supported",
        )

    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Uploaded file is empty",
        )

    content = extract_resume_text(filename, file_bytes)
    return save_resume_version(user_id=user_id, filename=filename, content=content, db=db)


@router.post(
    "/resumes/versions",
    response_model=ResumeVersionResponse,
    status_code=status.HTTP_201_CREATED,
)
def create_resume_version(
    resume_data: ResumeVersionCreate, db: Session = Depends(get_db)
):
    ensure_user_exists(resume_data.user_id, db)
    return save_resume_version(
        user_id=resume_data.user_id,
        filename=resume_data.filename,
        content=resume_data.content,
        db=db,
    )


@router.get("/resumes/versions", response_model=list[ResumeVersionResponse])
def list_resume_versions(user_id: int, db: Session = Depends(get_db)):
    return (
        db.query(ResumeVersion)
        .filter(ResumeVersion.user_id == user_id)
        .order_by(ResumeVersion.created_at.desc(), ResumeVersion.id.desc())
        .all()
    )


@router.post("/resumes/match-analyses", response_model=ResumeMatchAnalysisResponse)
def create_resume_match_analysis(
    analysis_data: ResumeMatchAnalysisCreate, db: Session = Depends(get_db)
):
    resume = (
        db.query(ResumeVersion)
        .filter(
            ResumeVersion.id == analysis_data.resume_version_id,
            ResumeVersion.user_id == analysis_data.user_id,
        )
        .first()
    )
    if not resume:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Resume version not found",
        )
    try:
        result = analyze_resume_match_with_llm(
            resume_content=resume.content,
            job_description=analysis_data.job_description,
        )
    except Exception as error:
        logger.exception("Resume JD match LLM analysis failed: resume_version_id=%s", resume.id)
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail=str(error),
        ) from error
    record = ResumeAnalysisRecord(
        user_id=analysis_data.user_id,
        resume_version_id=resume.id,
        job_title=analysis_data.job_title.strip() or "Target role",
        job_description=analysis_data.job_description.strip(),
        resume_score=result["resume_score"],
        match_score=result["match_score"],
        dimension_scores=result["dimension_scores"],
        resume_suggestions=result["resume_suggestions"],
        match_highlights=result["match_highlights"],
        match_gaps=result["match_gaps"],
    )
    db.add(record)
    db.commit()
    db.refresh(record)
    return record


@router.get("/resumes/match-analyses", response_model=list[ResumeMatchAnalysisResponse])
def list_resume_match_analyses(user_id: int, db: Session = Depends(get_db)):
    return (
        db.query(ResumeAnalysisRecord)
        .filter(ResumeAnalysisRecord.user_id == user_id)
        .order_by(ResumeAnalysisRecord.created_at.desc(), ResumeAnalysisRecord.id.desc())
        .all()
    )


@router.post("/projects/analyze", response_model=ProjectAnalysisResponse)
def analyze_project(project_data: ProjectAnalysisRequest, db: Session = Depends(get_db)):
    logger.info("Project analysis API called: project=%s", project_data.project_name)
    try:
        result = analyze_project_with_llm(project_data.model_dump())
        logger.info(
            "Project analysis API completed: project=%s model=%s matched=%s",
            project_data.project_name,
            result.get("model"),
            len(result.get("matched_capabilities", [])),
        )

        project = (
            db.query(ProjectCapabilityProfile)
            .filter(
                ProjectCapabilityProfile.user_id == project_data.user_id,
                func.lower(ProjectCapabilityProfile.project_name) == project_data.project_name.lower(),
            )
            .first()
        )

        dimension_scores = build_dimension_score_items(result)

        record = ProjectDimensionScoreRecord(
            project_id=project.id if project else 0,
            project_name=project_data.project_name,
            user_id=project_data.user_id,
            technology_stack=project_data.technology_stack,
            problem=project_data.problem,
            solution_and_validation=project_data.solution_and_validation,
            dimension_scores=dimension_scores,
            model=result.get("model") or "",
        )
        db.add(record)
        db.commit()
        db.refresh(record)
        logger.info("Project dimension score saved: project=%s record_id=%s", project_data.project_name, record.id)

        return result
    except Exception as error:
        logger.exception("Project analysis API failed: project=%s", project_data.project_name)
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail=str(error),
        )


@router.post("/resumes/analyze", response_model=ResumeAnalysisResponse)
def analyze_resume(resume_data: ResumeAnalysisRequest):
    logger.info("Resume analysis API called: filename=%s", resume_data.filename)
    try:
        result = analyze_resume_with_llm(resume_data.model_dump())
        logger.info(
            "Resume analysis API completed: filename=%s model=%s",
            resume_data.filename,
            result.get("model"),
        )
        return result
    except Exception as error:
        logger.exception("Resume analysis API failed: filename=%s", resume_data.filename)
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail=str(error),
        )


def build_dimension_score_items(result: dict) -> list[dict]:
    return [
        {
            "name": "Python编程与网络基础",
            "score": round(float(result.get("python_networking_score") or 0), 1),
            "suggestion": result.get("python_networking_reason") or "",
        },
        {
            "name": "大模型基础与Prompt工程",
            "score": round(float(result.get("llm_prompt_engineering_score") or 0), 1),
            "suggestion": result.get("llm_prompt_engineering_reason") or "",
        },
        {
            "name": "RAG检索增强技术",
            "score": round(float(result.get("rag_score") or 0), 1),
            "suggestion": result.get("rag_reason") or "",
        },
        {
            "name": "AI框架、智能体与多模态集成",
            "score": round(float(result.get("ai_frameworks_score") or 0), 1),
            "suggestion": result.get("ai_frameworks_reason") or "",
        },
        {
            "name": "后端服务与数据存储",
            "score": round(float(result.get("backend_data_storage_score") or 0), 1),
            "suggestion": result.get("backend_data_storage_reason") or "",
        },
        {
            "name": "模型部署、推理与微调",
            "score": round(float(result.get("deployment_finetuning_score") or 0), 1),
            "suggestion": result.get("deployment_finetuning_reason") or "",
        },
        {
            "name": "工程运维与业务交付",
            "score": round(float(result.get("engineering_ops_score") or 0), 1),
            "suggestion": result.get("engineering_ops_reason") or "",
        },
    ]


@router.get(
    "/resumes/dimension-scores",
    response_model=ResumeDimensionScoreRecordResponse | None,
)
def get_resume_dimension_score(
    user_id: int,
    resume_version_id: int,
    db: Session = Depends(get_db),
):
    return (
        db.query(ResumeDimensionScoreRecord)
        .filter(
            ResumeDimensionScoreRecord.user_id == user_id,
            ResumeDimensionScoreRecord.resume_version_id == resume_version_id,
        )
        .first()
    )


@router.get(
    "/projects/dimension-scores/total",
    response_model=dict,
)
def get_project_dimension_scores_total(
    user_id: int,
    db: Session = Depends(get_db),
):
    records = (
        db.query(ProjectDimensionScoreRecord)
        .filter(ProjectDimensionScoreRecord.user_id == user_id)
        .all()
    )
    
    total_scores = {}
    project_details = []
    
    for record in records:
        dimension_scores = record.dimension_scores if isinstance(record.dimension_scores, list) else []
        project_item = {
            "project_name": record.project_name,
            "created_at": record.created_at.isoformat() if record.created_at else None,
            "scores": []
        }
        
        for dim in dimension_scores:
            dim_name = dim.get("name", "")
            dim_score = dim.get("score", 0)
            
            if dim_name not in total_scores:
                total_scores[dim_name] = 0
            total_scores[dim_name] += dim_score
            
            project_item["scores"].append({
                "name": dim_name,
                "score": dim_score
            })
        
        project_details.append(project_item)
    
    return {
        "total_scores": total_scores,
        "project_details": project_details
    }


@router.post(
    "/resumes/analyze-dimensions",
    response_model=ResumeDimensionScoreRecordResponse,
)
def analyze_resume_dimensions(
    resume_data: ResumeDimensionsRequest,
    db: Session = Depends(get_db),
):
    logger.info(
        "Resume dimensions analysis API called: resume_version_id=%s",
        resume_data.resume_version_id,
    )
    try:
        resume = (
            db.query(ResumeVersion)
            .filter(
                ResumeVersion.id == resume_data.resume_version_id,
                ResumeVersion.user_id == resume_data.user_id,
            )
            .first()
        )
        if not resume:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Resume version not found",
            )

        result = analyze_resume_dimensions_with_llm(
            {"filename": resume.filename, "content": resume.content}
        )
        dimension_scores = build_dimension_score_items(result)
        resume_score = round(
            sum(item["score"] for item in dimension_scores) / len(dimension_scores)
        )
        record = (
            db.query(ResumeDimensionScoreRecord)
            .filter(
                ResumeDimensionScoreRecord.user_id == resume_data.user_id,
                ResumeDimensionScoreRecord.resume_version_id == resume_data.resume_version_id,
            )
            .first()
        )
        if record is None:
            record = ResumeDimensionScoreRecord(
                user_id=resume_data.user_id,
                resume_version_id=resume_data.resume_version_id,
            )
            db.add(record)
        record.resume_score = resume_score
        record.dimension_scores = dimension_scores
        record.model = result.get("model") or ""
        db.commit()
        db.refresh(record)
        logger.info(
            "Resume dimensions analysis API completed: resume_version_id=%s model=%s",
            resume_data.resume_version_id,
            record.model,
        )
        return record
    except Exception as error:
        if isinstance(error, HTTPException):
            raise
        logger.exception(
            "Resume dimensions analysis API failed: resume_version_id=%s",
            resume_data.resume_version_id,
        )
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail=str(error),
        )


@router.get(
    "/projects/dimension-scores",
    response_model=list[ProjectDimensionScoreRecordResponse],
)
def get_project_dimension_scores(
    user_id: int,
    project_id: int | None = None,
    db: Session = Depends(get_db),
):
    query = db.query(ProjectDimensionScoreRecord).filter(
        ProjectDimensionScoreRecord.user_id == user_id
    )
    if project_id is not None:
        query = query.filter(ProjectDimensionScoreRecord.project_id == project_id)
    return query.order_by(ProjectDimensionScoreRecord.created_at.desc()).all()


@router.post(
    "/projects/analyze-dimensions",
    response_model=ProjectDimensionScoreRecordResponse,
)
def analyze_project_dimensions(
    project_data: ProjectDimensionScoreRequest,
    db: Session = Depends(get_db),
):
    logger.info(
        "Project dimensions analysis API called: project_id=%s",
        project_data.project_id,
    )
    try:
        project = (
            db.query(ProjectCapabilityProfile)
            .filter(
                ProjectCapabilityProfile.id == project_data.project_id,
                ProjectCapabilityProfile.user_id == project_data.user_id,
            )
            .first()
        )
        if not project:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Project not found",
            )

        result = analyze_project_with_llm(
            {
                "project_name": project.project_name,
                "technology_stack": project.technology_stack,
                "problem": project.problem,
                "solution_and_validation": project.solution_and_validation,
            }
        )
        dimension_scores = build_dimension_score_items(result)

        record = ProjectDimensionScoreRecord(
            project_id=project.id,
            project_name=project.project_name,
            user_id=project.user_id,
            technology_stack=project.technology_stack,
            problem=project.problem,
            solution_and_validation=project.solution_and_validation,
            dimension_scores=dimension_scores,
            model=result.get("model") or "",
        )
        db.add(record)
        db.commit()
        db.refresh(record)
        logger.info(
            "Project dimensions analysis API completed: project_id=%s model=%s",
            project_data.project_id,
            record.model,
        )
        return record
    except Exception as error:
        if isinstance(error, HTTPException):
            raise
        logger.exception(
            "Project dimensions analysis API failed: project_id=%s",
            project_data.project_id,
        )
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail=str(error),
        )


@router.get(
    "/learning-backlog",
    response_model=list[LearningBacklogResponse],
)
def list_learning_backlog(user_id: int, db: Session = Depends(get_db)):
    return (
        db.query(LearningBacklog)
        .filter(LearningBacklog.user_id == user_id)
        .order_by(LearningBacklog.added_at.desc(), LearningBacklog.id.desc())
        .all()
    )


@router.post(
    "/learning-backlog",
    response_model=LearningBacklogResponse,
    status_code=status.HTTP_201_CREATED,
)
def create_learning_backlog(
    item_data: LearningBacklogCreate, db: Session = Depends(get_db)
):
    ensure_user_exists(item_data.user_id, db)
    title = item_data.title.strip()
    if not title:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Learning backlog title cannot be empty",
        )
    item = LearningBacklog(
        user_id=item_data.user_id,
        title=title,
        completed=False,
    )
    if item_data.added_at is not None:
        item.added_at = item_data.added_at
    db.add(item)
    db.commit()
    db.refresh(item)
    return item


@router.patch(
    "/learning-backlog/{item_id}/toggle",
    response_model=LearningBacklogResponse,
)
def toggle_learning_backlog(
    item_id: int, user_id: int, db: Session = Depends(get_db)
):
    item = (
        db.query(LearningBacklog)
        .filter(
            LearningBacklog.id == item_id,
            LearningBacklog.user_id == user_id,
        )
        .first()
    )
    if not item:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Learning backlog item not found",
        )
    item.completed = not item.completed
    item.completed_at = datetime.now(timezone.utc) if item.completed else None
    db.commit()
    db.refresh(item)
    return item


@router.delete(
    "/learning-backlog/{item_id}",
    status_code=status.HTTP_204_NO_CONTENT,
)
def delete_learning_backlog(
    item_id: int, user_id: int, db: Session = Depends(get_db)
):
    item = (
        db.query(LearningBacklog)
        .filter(
            LearningBacklog.id == item_id,
            LearningBacklog.user_id == user_id,
        )
        .first()
    )
    if not item:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Learning backlog item not found",
        )
    db.delete(item)
    db.commit()
    return None
